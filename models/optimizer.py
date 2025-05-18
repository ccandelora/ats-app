import re
from .analyzer import ResumeAnalyzer
from docx import Document
import docx
import io

class ResumeOptimizer:
    def __init__(self):
        self.analyzer = ResumeAnalyzer()
        # Add statistics tracking
        self.stats = {
            'keywords_added': 0,
            'action_verbs_enhanced': 0,
            'quantifiable_added': 0,
            'format_issues_fixed': 0
        }
        
    def get_optimization_stats(self):
        """Return optimization statistics."""
        return self.stats

    def optimize_resume(self, resume_text, job_description, industry='general', customizations=None):
        """Generate an optimized resume based on job description and industry."""
        # Reset statistics for new optimization
        self.stats = {
            'keywords_added': 0,
            'action_verbs_enhanced': 0,
            'quantifiable_added': 0,
            'format_issues_fixed': 0
        }
        
        # Process customizations
        if customizations is None:
            customizations = {}
        
        # Apply industry-specific optimization techniques
        industry_settings = self._get_industry_settings(industry)
        
        # Analyze the original resume
        analysis = self.analyzer.analyze_resume(resume_text, job_description)
        
        # Extract missing keywords from job description
        missing_keywords = analysis["match_results"]["missing_keywords"]
        
        # Prioritize industry-specific keywords if available
        if industry != 'general':
            # Use industry-specific weights to sort keywords
            from .industry_scorer import IndustryScorer
            industry_scorer = IndustryScorer()
            
            industry_keywords = industry_scorer.INDUSTRY_KEYWORD_WEIGHTS.get(industry, {})
            
            # Sort missing keywords by industry importance
            sorted_missing = []
            for keyword in missing_keywords:
                weight = industry_keywords.get(keyword.lower(), 1.0)
                sorted_missing.append((keyword, weight))
            
            # Sort by weight descending
            sorted_missing.sort(key=lambda x: x[1], reverse=True)
            missing_keywords = [kw for kw, _ in sorted_missing]
        
        # Identify resume sections
        sections = analysis["sections"]
        
        # Initialize optimized resume with original content
        optimized_sections = {}
        for section_name, section_data in sections.items():
            if "content" in section_data:
                optimized_sections[section_name] = section_data["content"]
        
        # If we couldn't identify sections properly, use the whole text
        if not optimized_sections:
            # Just add missing keywords to the summary/profile section
            lines = resume_text.split('\n')
            # Look for a summary/profile section
            summary_index = -1
            for i, line in enumerate(lines):
                if re.search(r'(?i)(summary|profile|objective)', line):
                    summary_index = i
                    break
            
            if summary_index >= 0 and summary_index + 1 < len(lines):
                # Add optimized content after the summary heading
                summary_text = lines[summary_index + 1]
                enriched_summary = self._enrich_text(summary_text, missing_keywords, industry_settings)
                lines[summary_index + 1] = enriched_summary
                return '\n'.join(lines)
            else:
                # No identifiable sections, return with minimal changes
                return resume_text
                
        # Enhance summary/profile section
        if "summary" in optimized_sections:
            optimized_sections["summary"] = self._enrich_text(
                optimized_sections["summary"], 
                missing_keywords[:5],  # Add top 5 missing keywords to summary
                industry_settings
            )
            
        # Enhance skills section
        if "skills" in optimized_sections:
            # Add missing keywords as skills
            skills_text = optimized_sections["skills"]
            skills_lines = skills_text.split('\n')
            
            # Find where to insert new skills
            if len(skills_lines) > 0:
                new_skills = []
                added_keywords = set()
                
                # Add industry-specific skills first
                for keyword in missing_keywords:
                    # Avoid duplicates and only add relevant skills
                    if keyword.lower() not in skills_text.lower() and keyword not in added_keywords and len(keyword) > 3:
                        # Format based on existing skill list format
                        if any(line.strip().startswith('•') for line in skills_lines):
                            new_skills.append(f"• {keyword.title()}")
                        elif any(line.strip().startswith('-') for line in skills_lines):
                            new_skills.append(f"- {keyword.title()}")
                        else:
                            new_skills.append(f"{keyword.title()}")
                        added_keywords.add(keyword)
                
                if new_skills:
                    # Find a good insertion point (after existing bullet points)
                    insertion_point = 0
                    for i, line in enumerate(skills_lines):
                        if line.strip().startswith('•') or line.strip().startswith('-'):
                            insertion_point = i + 1
                    
                    # Insert new skills
                    skills_lines = skills_lines[:insertion_point] + new_skills + skills_lines[insertion_point:]
                    optimized_sections["skills"] = '\n'.join(skills_lines)
        
        # Enhance experience section to include relevant keywords
        if "experience" in optimized_sections:
            experience_text = optimized_sections["experience"]
            # Look for places to naturally integrate keywords
            experience_lines = experience_text.split('\n')
            
            # Process bullet points to include keywords
            for i in range(len(experience_lines)):
                line = experience_lines[i].strip()
                if (line.startswith('•') or line.startswith('-')) and len(line) > 10:
                    # This is a bullet point - use industry-specific enhancement
                    experience_lines[i] = self._enhance_bullet_point(
                        line, 
                        missing_keywords,
                        industry_settings
                    )
            
            optimized_sections["experience"] = '\n'.join(experience_lines)
        
        # Add action verbs if needed
        if "experience" in optimized_sections and not analysis["action_verbs"]["has_action_verbs"]:
            optimized_sections["experience"] = self._add_action_verbs(
                optimized_sections["experience"], 
                analysis["action_verbs"]["missing"],
                industry_settings.get('preferred_verbs', [])
            )
        
        # Add quantifiable results if needed
        if "experience" in optimized_sections and not analysis["quantifiable_results"]["has_quantifiable_results"]:
            optimized_sections["experience"] = self._suggest_quantifiable_results(
                optimized_sections["experience"],
                industry_settings.get('metrics', [])
            )
        
        # Reconstruct the resume with customizations
        reconstructed_resume = []
        
        # Add custom contact information if provided
        contact_info = customizations.get('contact_info')
        if contact_info and any(contact_info.values()):
            # Format contact info at the top of the resume
            contact_block = []
            if contact_info.get('name'):
                contact_block.append(contact_info['name'])
            
            contact_details = []
            if contact_info.get('email'):
                contact_details.append(contact_info['email'])
            if contact_info.get('phone'):
                contact_details.append(contact_info['phone'])
            if contact_info.get('linkedin'):
                contact_details.append(contact_info['linkedin'])
            
            if contact_details:
                contact_block.append(' | '.join(contact_details))
            
            reconstructed_resume.extend(contact_block)
            reconstructed_resume.append('')  # Add blank line after contact info
        
        # Add sections in their original order
        for section_name, section_content in sections.items():
            if "start" in section_content and "end" in section_content:
                # Get the original section header
                original_lines = resume_text.split('\n')
                header_line = original_lines[section_content["start"]]
                reconstructed_resume.append(header_line)
                
                # Add optimized content if available, otherwise original content
                if section_name in optimized_sections:
                    reconstructed_resume.append(optimized_sections[section_name])
                else:
                    section_lines = original_lines[section_content["start"]+1:section_content["end"]]
                    reconstructed_resume.append('\n'.join(section_lines))
                
                reconstructed_resume.append('')  # Add blank line between sections
        
        return '\n'.join(reconstructed_resume)
    
    def _get_industry_settings(self, industry):
        """Get industry-specific optimization settings."""
        # Default settings
        settings = {
            'preferred_verbs': [
                'achieved', 'improved', 'led', 'managed', 'developed'
            ],
            'metrics': [
                'by X%', 'for X clients', 'across X teams', 'resulting in $X'
            ],
            'keyword_emphasis': 'moderate'  # options: light, moderate, strong
        }
        
        # Industry-specific settings
        if industry == 'tech':
            settings.update({
                'preferred_verbs': [
                    'developed', 'implemented', 'architected', 'programmed', 'engineered',
                    'designed', 'optimized', 'deployed', 'debugged', 'maintained'
                ],
                'metrics': [
                    'reducing latency by X%', 'improving efficiency by X%', 
                    'for X users', 'processing X transactions per second',
                    'saving $X in operational costs'
                ],
                'keyword_emphasis': 'strong'
            })
        elif industry == 'finance':
            settings.update({
                'preferred_verbs': [
                    'analyzed', 'forecasted', 'reconciled', 'projected', 'audited',
                    'budgeted', 'allocated', 'evaluated', 'assessed', 'monitored'
                ],
                'metrics': [
                    'managing $X in assets', 'increasing revenue by X%', 
                    'reducing costs by $X', 'for X accounts',
                    'improving accuracy by X%'
                ],
                'keyword_emphasis': 'moderate'
            })
        elif industry == 'healthcare':
            settings.update({
                'preferred_verbs': [
                    'treated', 'diagnosed', 'administered', 'assessed', 'monitored',
                    'coordinated', 'provided', 'documented', 'educated', 'implemented'
                ],
                'metrics': [
                    'for X patients', 'improving outcomes by X%', 
                    'reducing readmissions by X%', 'with X compliance rate',
                    'saving X hours of staff time'
                ],
                'keyword_emphasis': 'moderate'
            })
        elif industry == 'marketing':
            settings.update({
                'preferred_verbs': [
                    'launched', 'promoted', 'drove', 'created', 'designed',
                    'executed', 'generated', 'strategized', 'coordinated', 'analyzed'
                ],
                'metrics': [
                    'increasing conversions by X%', 'generating X leads', 
                    'achieving X% ROI', 'for X campaigns',
                    'growing audience by X%'
                ],
                'keyword_emphasis': 'strong'
            })
        elif industry == 'sales':
            settings.update({
                'preferred_verbs': [
                    'sold', 'negotiated', 'acquired', 'closed', 'generated',
                    'exceeded', 'cultivated', 'secured', 'prospected', 'presented'
                ],
                'metrics': [
                    'exceeding quota by X%', 'generating $X in revenue', 
                    'closing X deals', 'retaining X% of clients',
                    'growing territory by X%'
                ],
                'keyword_emphasis': 'strong'
            })
        
        return settings
    
    def _enrich_text(self, text, keywords, industry_settings, max_keywords=5):
        """Add keywords to text where appropriate."""
        if not text or not keywords:
            return text
            
        # Determine how many keywords to add based on industry emphasis
        emphasis = industry_settings.get('keyword_emphasis', 'moderate')
        if emphasis == 'light':
            max_keywords = min(3, max_keywords)
        elif emphasis == 'strong':
            max_keywords = max_keywords + 1
            
        # Simple approach: add keywords to the end of the text
        keywords_to_add = [k for k in keywords[:max_keywords] if k.lower() not in text.lower()]
        
        # Update statistics
        if keywords_to_add:
            self.stats['keywords_added'] += len(keywords_to_add)
        
        if not keywords_to_add:
            return text
            
        enriched_text = text.rstrip()
        
        # Check if the last sentence ends with proper punctuation
        if not enriched_text.endswith('.') and not enriched_text.endswith('!') and not enriched_text.endswith('?'):
            enriched_text += '.'
            
        # Add keywords in a natural way
        keyword_phrase = ", ".join(k.title() for k in keywords_to_add)
        
        # Vary the phrasing based on industry
        phrases = [
            f" Experienced in {keyword_phrase}.",
            f" Proficient with {keyword_phrase}.",
            f" Skilled in {keyword_phrase}.",
            f" Knowledgeable about {keyword_phrase}.",
            f" Well-versed in {keyword_phrase}."
        ]
        
        # Pick a phrase based on the first character of the text for consistency
        phrase_index = ord(text[0]) % len(phrases) if text else 0
        enriched_text += phrases[phrase_index]
            
        return enriched_text
    
    def _enhance_bullet_point(self, bullet_point, keywords, industry_settings, max_keywords=2):
        """Enhance a bullet point with relevant keywords."""
        # Only process if there are keywords to add
        if not keywords:
            return bullet_point
            
        # Adjust max keywords based on industry emphasis
        emphasis = industry_settings.get('keyword_emphasis', 'moderate')
        if emphasis == 'light':
            max_keywords = 1
        elif emphasis == 'strong':
            max_keywords = 3
            
        # Find keywords that aren't already in the bullet point
        keywords_to_add = []
        for keyword in keywords:
            if keyword.lower() not in bullet_point.lower() and len(keywords_to_add) < max_keywords:
                keywords_to_add.append(keyword)
                
        # Update statistics
        if keywords_to_add:
            self.stats['keywords_added'] += len(keywords_to_add)
                
        if not keywords_to_add:
            return bullet_point
            
        # Add keywords to the bullet point in a natural way
        enhanced_bullet = bullet_point.rstrip()
        
        # Check if the bullet point ends with proper punctuation
        if not enhanced_bullet.endswith('.') and not enhanced_bullet.endswith('!') and not enhanced_bullet.endswith('?'):
            enhanced_bullet += '.'
            
        # Add a phrase with the keyword
        keyword_phrase = " and ".join(k.lower() for k in keywords_to_add)
        
        # Choose a connecting phrase based on the bullet content and industry
        if "developed" in enhanced_bullet.lower() or "created" in enhanced_bullet.lower():
            enhanced_bullet = enhanced_bullet[:-1] + f", utilizing {keyword_phrase}."
        elif "managed" in enhanced_bullet.lower() or "led" in enhanced_bullet.lower():
            enhanced_bullet = enhanced_bullet[:-1] + f", with focus on {keyword_phrase}."
        elif "analyzed" in enhanced_bullet.lower() or "researched" in enhanced_bullet.lower():
            enhanced_bullet = enhanced_bullet[:-1] + f", particularly in relation to {keyword_phrase}."
        else:
            enhanced_bullet = enhanced_bullet[:-1] + f", incorporating {keyword_phrase}."
            
        return enhanced_bullet
    
    def _add_action_verbs(self, experience_text, missing_verbs, industry_verbs=None):
        """Add action verbs to experience bullet points."""
        if not missing_verbs and not industry_verbs:
            return experience_text
            
        # Prioritize industry-specific verbs if available
        verbs_to_use = industry_verbs or missing_verbs
        
        lines = experience_text.split('\n')
        verb_index = 0
        verbs_added = 0
        
        for i in range(len(lines)):
            line = lines[i].strip()
            # Only modify bullet points that don't start with action verbs
            if (line.startswith('•') or line.startswith('-')) and not any(line.lower().startswith(f"• {v}") or line.lower().startswith(f"- {v}") for v in verbs_to_use):
                # Strip the bullet and leading whitespace
                content = line[1:].strip()
                
                # Add the action verb at the beginning
                verb = verbs_to_use[verb_index % len(verbs_to_use)].title()
                lines[i] = f"{line[0]} {verb} {content}"
                verb_index += 1
                verbs_added += 1
        
        # Update statistics
        self.stats['action_verbs_enhanced'] += verbs_added
        
        return '\n'.join(lines)
    
    def _suggest_quantifiable_results(self, experience_text, industry_metrics=None):
        """Suggest adding quantifiable results to experience bullet points."""
        lines = experience_text.split('\n')
        modified = False
        
        # Default metrics if not provided
        metrics = industry_metrics or [
            "resulting in X% improvement",
            "saving $X annually",
            "increasing efficiency by X%",
            "for X clients/customers",
            "reducing costs by X%"
        ]
        
        for i in range(len(lines)):
            line = lines[i].strip()
            # Look for bullet points without numbers
            if (line.startswith('•') or line.startswith('-')) and not re.search(r'\d+', line) and len(line) > 15:
                # Choose appropriate metric based on content
                metric_index = i % len(metrics)
                if "cost" in line.lower() or "budget" in line.lower():
                    metric = "reducing costs by X%"
                elif "sales" in line.lower() or "revenue" in line.lower():
                    metric = "increasing revenue by X%"
                elif "team" in line.lower() or "staff" in line.lower():
                    metric = "leading a team of X people"
                else:
                    metric = metrics[metric_index]
                
                # Add a placeholder for quantifiable result
                lines[i] = f"{line} [{metric}]"
                modified = True
                
                # Update statistics
                self.stats['quantifiable_added'] += 1
                
                break  # Only modify one line as an example
        
        return '\n'.join(lines)
        
    def generate_docx(self, resume_text):
        """Generate a professionally formatted DOCX document from the optimized resume text."""
        doc = Document()
        
        # Set document styles
        styles = doc.styles
        style = styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = docx.shared.Pt(11)
        
        # Extract contact info from the resume for the header
        from utils.parser import extract_contact_info
        contact_info = extract_contact_info(resume_text)
        
        # Create professional header with contact information
        if contact_info['name']:
            # Add name as large, bold header
            name_paragraph = doc.add_paragraph()
            name_run = name_paragraph.add_run(contact_info['name'])
            name_run.bold = True
            name_run.font.size = docx.shared.Pt(16)
            name_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            # Add contact details centered under name
            contact_details = []
            if contact_info['email']:
                contact_details.append(contact_info['email'])
            if contact_info['phone']:
                contact_details.append(contact_info['phone'])
            if contact_info['linkedin']:
                contact_details.append(contact_info['linkedin'])
                
            if contact_details:
                contact_paragraph = doc.add_paragraph()
                contact_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                contact_paragraph.add_run(" | ".join(contact_details))
                
            # Add a separator line
            doc.add_paragraph()
        
        # Process resume sections
        lines = resume_text.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip lines with contact info since we've already added it at the top
            if (contact_info['name'] and line == contact_info['name']) or \
               (contact_info['email'] and line == contact_info['email']) or \
               (contact_info['phone'] and line == contact_info['phone']) or \
               (contact_info['linkedin'] and line == contact_info['linkedin']):
                i += 1
                continue
                
            # Check if this is a section header
            if line and (line.isupper() or any(section in line.lower() for section in ['summary', 'experience', 'education', 'skills'])):
                # Add section header with appropriate formatting
                header_paragraph = doc.add_paragraph()
                header_run = header_paragraph.add_run(line)
                header_run.bold = True
                header_run.font.size = docx.shared.Pt(14)
                
                # Add a separator line under the header
                doc.add_paragraph()
                i += 1
                
                # Process the section content
                while i < len(lines) and (not lines[i].strip() or not (lines[i].strip().isupper() or 
                       any(section in lines[i].lower() for section in ['summary', 'experience', 'education', 'skills']))):
                    line = lines[i].strip()
                    
                    if not line:
                        i += 1
                        continue
                        
                    if line.startswith('•') or line.startswith('-'):
                        # This is a bullet point
                        paragraph = doc.add_paragraph()
                        paragraph.style = 'List Bullet'
                        paragraph.add_run(line[1:].strip())
                    else:
                        # Regular paragraph
                        paragraph = doc.add_paragraph()
                        paragraph.add_run(line)
                        
                    i += 1
            else:
                # Regular paragraph
                if line:
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(line)
                i += 1
        
        # Save to memory
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream.getvalue() 